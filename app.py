"""
HackRX 5.0 - LLM-Powered Intelligent Query-Retrieval System
FastAPI application for processing documents and answering questions
"""

import os
import json
import asyncio
from datetime import datetime
from typing import List, Dict, Any, Optional
import logging
from pathlib import Path

# FastAPI imports
from fastapi import FastAPI, HTTPException, Depends, status
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field, validator
import uvicorn

# Document processing imports
import requests
from docx import Document
import email
from email import policy
from email.parser import BytesParser
import tempfile
import shutil
from urllib.parse import urlparse

# LangChain and ML imports
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.prompts import PromptTemplate
from langchain.llms.base import LLM
from langchain.memory import ConversationBufferMemory
from groq import Groq
import torch
import gc

# Import configuration
from config import Config

# Configure logging
logging.basicConfig(level=getattr(logging, Config.LOG_LEVEL))
logger = logging.getLogger(__name__)

# Validate configuration on startup
Config.validate_config()

# Initialize FastAPI app
app = FastAPI(
    title="HackRX 5.0 - Intelligent Query-Retrieval System",
    description="LLM-Powered system for processing documents and answering contextual queries",
    version="1.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer()

# Groq client configuration
groq_client = Groq(api_key=Config.GROQ_API_KEY)

# Global variables
vectorstore = None
qa_chain = None

# Pydantic models
class QueryRequest(BaseModel):
    documents: str = Field(..., description="URL to the document blob")
    questions: List[str] = Field(..., min_items=1, description="List of questions to answer")
    
    @validator('documents')
    def validate_documents_url(cls, v):
        if not v.startswith(('http://', 'https://')):
            raise ValueError('Documents must be a valid URL')
        return v

class QueryResponse(BaseModel):
    answers: List[str] = Field(..., description="List of answers corresponding to the questions")
    metadata: Optional[Dict[str, Any]] = Field(default=None, description="Additional metadata")

# Custom Groq LLM class optimized for concise responses
class GroqLLM(LLM):
    model_name: str = "llama3-8b-8192"  # Fast model for low latency
    client: Any = Field(default=None)
    temperature: float = 0.1  # Low temperature for consistent, fast responses
    max_tokens: int = 200  # Reduced for concise responses

    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.client = groq_client

    @property
    def _llm_type(self) -> str:
        return "groq"

    def _call(self, prompt: str, stop: Optional[List[str]] = None, run_manager: Optional[Any] = None) -> str:
        try:
            # Optimize prompt for very concise, direct response
            optimized_prompt = f"{prompt}\n\nAnswer in maximum 25 words. Be direct and include only key facts:"
            
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=[
                    {"role": "user", "content": optimized_prompt}
                ],
                temperature=self.temperature,
                max_tokens=self.max_tokens,
                stream=False  # Disable streaming for faster processing in batch
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Error in Groq API call: {str(e)}")
            return "I apologize, but I encountered an error processing your request."

    @property
    def _identifying_params(self) -> Dict[str, Any]:
        return {"model_name": self.model_name, "temperature": self.temperature}

# Document processing utilities
class DocumentProcessor:
    @staticmethod
    async def download_document(url: str) -> str:
        """Download document from URL and return local path"""
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Create temporary file with proper extension extraction
            temp_dir = tempfile.mkdtemp()
            
            # Extract file extension from URL path, ignoring query parameters
            parsed_url = urlparse(url)
            url_path = parsed_url.path
            
            # Get file extension from the path
            file_extension = Path(url_path).suffix or '.pdf'
            
            # Create safe filename
            temp_file = os.path.join(temp_dir, f"document{file_extension}")
            
            with open(temp_file, 'wb') as f:
                f.write(response.content)
            
            logger.info(f"Downloaded document to: {temp_file}")
            return temp_file
            
        except Exception as e:
            logger.error(f"Error downloading document: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to download document: {str(e)}")

    @staticmethod
    def load_pdf(file_path: str) -> List[Any]:
        """Load PDF document"""
        try:
            loader = PyPDFLoader(file_path)
            documents = loader.load()
            logger.info(f"Loaded PDF with {len(documents)} pages")
            return documents
        except Exception as e:
            logger.error(f"Error loading PDF: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to load PDF: {str(e)}")

    @staticmethod
    def load_docx(file_path: str) -> List[Any]:
        """Load DOCX document"""
        try:
            from langchain.schema import Document as LangchainDocument
            
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text.append(paragraph.text)
            
            full_text = '\n\n'.join(text)
            documents = [LangchainDocument(page_content=full_text, metadata={"source": file_path})]
            logger.info(f"Loaded DOCX with {len(documents)} documents")
            return documents
        except Exception as e:
            logger.error(f"Error loading DOCX: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to load DOCX: {str(e)}")

    @staticmethod
    def load_email(file_path: str) -> List[Any]:
        """Load email document"""
        try:
            from langchain.schema import Document as LangchainDocument
            
            with open(file_path, 'rb') as f:
                msg = BytesParser(policy=policy.default).parse(f)
            
            # Extract email content
            subject = msg['subject'] or "No Subject"
            sender = msg['from'] or "Unknown Sender"
            body = ""
            
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body += part.get_content()
            else:
                body = msg.get_content()
            
            content = f"Subject: {subject}\nFrom: {sender}\n\n{body}"
            documents = [LangchainDocument(page_content=content, metadata={"source": file_path})]
            logger.info(f"Loaded email document")
            return documents
        except Exception as e:
            logger.error(f"Error loading email: {str(e)}")
            raise HTTPException(status_code=400, detail=f"Failed to load email: {str(e)}")

    @classmethod
    async def process_document(cls, url: str) -> List[Any]:
        """Process document based on file type"""
        file_path = await cls.download_document(url)
        
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return cls.load_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return cls.load_docx(file_path)
            elif file_extension in ['.eml', '.msg']:
                return cls.load_email(file_path)
            else:
                # Try to load as PDF by default
                return cls.load_pdf(file_path)
                
        finally:
            # Cleanup temporary file
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    temp_dir = os.path.dirname(file_path)
                    if os.path.exists(temp_dir):
                        shutil.rmtree(temp_dir)
            except Exception as e:
                logger.warning(f"Failed to cleanup temporary files: {str(e)}")

# Vector store and QA chain setup
class IntelligentQASystem:
    def __init__(self):
        self.device = 'cuda' if torch.cuda.is_available() else 'cpu'
        self.vectorstore = None
        self.qa_chain = None
        
    def create_embeddings(self, documents: List[Any]) -> FAISS:
        """Create FAISS vector store from documents"""
        try:
            logger.info("Creating embeddings...")
            
            # Split documents into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1500,
                chunk_overlap=250,
                length_function=len
            )
            texts = text_splitter.split_documents(documents)
            logger.info(f"Created {len(texts)} text chunks")
            
            # Create embeddings
            try:
                # Try using HuggingFaceEmbeddings first
                embeddings = HuggingFaceEmbeddings(
                    model_name="sentence-transformers/all-MiniLM-L6-v2",
                    model_kwargs={'device': self.device}
                )
                print("hug")
            except Exception as e:
                logger.warning(f"HuggingFaceEmbeddings failed: {e}, trying alternative...")
                # Fallback to sentence-transformers directly
                from sentence_transformers import SentenceTransformer
                from langchain_core.embeddings import Embeddings
                
                class SentenceTransformerEmbeddings(Embeddings):
                    def __init__(self, model_name: str):
                        self.model = SentenceTransformer(model_name)
                    
                    def embed_documents(self, texts: List[str]) -> List[List[float]]:
                        return self.model.encode(texts).tolist()
                    
                    def embed_query(self, text: str) -> List[float]:
                        return self.model.encode([text])[0].tolist()
                
                embeddings = SentenceTransformerEmbeddings("sentence-transformers/all-MiniLM-L6-v2")
            
            # Create FAISS vectorstore
            vectorstore = FAISS.from_documents(texts, embeddings)
            logger.info("Vector store created successfully")
            
            return vectorstore
            
        except Exception as e:
            logger.error(f"Error creating embeddings: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to create embeddings: {str(e)}")
    
    def setup_qa_chain(self, vectorstore: FAISS) -> ConversationalRetrievalChain:
        """Set up the conversational QA chain"""
        try:
            logger.info("Setting up QA chain...")
            
            llm = GroqLLM()
            
            # Enhanced prompt template for structured responses
            prompt_template = """You are an intelligent document analysis assistant specializing in insurance, legal, HR, and compliance domains.

Based on the provided context, answer the question with a single, comprehensive, and precise sentence that captures all essential information.

Context: {context}

Question: {question}

Instructions:
1. Provide a complete, well-structured sentence that includes all relevant details
2. Include specific numbers, dates, conditions, and requirements when available
3. Be precise and factual, avoiding speculation
4. If information is not available in the context, state this clearly
5. Focus on the most important and actionable information

Answer:"""

            qa_chain = ConversationalRetrievalChain.from_llm(
                llm=llm,
                retriever=vectorstore.as_retriever(
                    search_type="similarity",
                    search_kwargs={'k': 5}
                ),
                memory=ConversationBufferMemory(
                    memory_key="chat_history",
                    return_messages=True,
                    output_key="answer"
                ),
                return_source_documents=True,
                combine_docs_chain_kwargs={"prompt": PromptTemplate(
                    template=prompt_template,
                    input_variables=["context", "question"]
                )},
                verbose=False
            )
            
            logger.info("QA chain setup completed")
            return qa_chain
            
        except Exception as e:
            logger.error(f"Error setting up QA chain: {str(e)}")
            raise HTTPException(status_code=500, detail=f"Failed to setup QA chain: {str(e)}")

# Authentication dependency
async def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if credentials.credentials != Config.BEARER_TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication credentials",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return credentials.credentials

# API endpoints
@app.get("/")
async def root():
    """Root endpoint"""
    return {
        "message": "HackRX 5.0 - Intelligent Query-Retrieval System",
        "version": "1.0.0",
        "status": "active"
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "device": 'cuda' if torch.cuda.is_available() else 'cpu'
    }

@app.post("/hackrx/run", response_model=QueryResponse)
async def run_hackrx(
    request: QueryRequest,
    token: str = Depends(verify_token)
) -> QueryResponse:
    """
    Main endpoint for processing documents and answering questions
    """
    try:
        logger.info(f"Processing request with {len(request.questions)} questions")
        
        # Initialize QA system
        qa_system = IntelligentQASystem()
        
        # Process document
        logger.info("Processing document...")
        documents = await DocumentProcessor.process_document(request.documents)
        
        if not documents:
            raise HTTPException(status_code=400, detail="No documents could be processed")
        
        # Create vector store
        logger.info("Creating vector store...")
        vectorstore = qa_system.create_embeddings(documents)
        
        # Setup QA chain
        logger.info("Setting up QA chain...")
        qa_chain = qa_system.setup_qa_chain(vectorstore)
        
        # Process questions
        logger.info("Processing questions...")
        answers = []
        
        for i, question in enumerate(request.questions):
            try:
                logger.info(f"Processing question {i+1}/{len(request.questions)}: {question[:50]}...")
                
                result = qa_chain({"question": question})
                answer = result["answer"].strip()
                
                # Ensure answer is a single comprehensive sentence
                if not answer.endswith('.'):
                    answer += '.'
                
                answers.append(answer)
                logger.info(f"Answer {i+1} generated successfully")
                
            except Exception as e:
                logger.error(f"Error processing question {i+1}: {str(e)}")
                answers.append("I apologize, but I could not process this question due to an error.")
        
        # Cleanup
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
        gc.collect()
        
        logger.info("Request processing completed successfully")
        
        return QueryResponse(
            answers=answers,
            metadata={
                "document_url": request.documents,
                "total_questions": len(request.questions),
                "processing_timestamp": datetime.now().isoformat(),
                "model_info": {
                    "llm": "groq-llama3-8b-8192",
                    "embeddings": "sentence-transformers/all-MiniLM-L6-v2",
                    "vectorstore": "FAISS"
                }
            }
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Internal server error: {str(e)}")

# Error handlers
from fastapi.responses import JSONResponse

@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "timestamp": datetime.now().isoformat()
        }
    )

@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    logger.error(f"Unhandled exception: {str(exc)}")
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "status_code": 500,
            "timestamp": datetime.now().isoformat()
        }
    )

if __name__ == "__main__":
    uvicorn.run(
        "app:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )
